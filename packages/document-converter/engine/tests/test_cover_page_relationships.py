"""Regression tests for cover-page relationship remapping (issue #506).

`copy_template_cover_page` deep-copies the template's body elements into the
pandoc-generated output. Those elements carry **template-local** `r:id`s — and
pandoc numbered the output's `word/_rels/document.xml.rels` independently. When
the two numberings disagree, the copied cover `<w:sectPr>` ends up pointing at
ids that mean something else, or at nothing at all (the reporter's dangling
`rId15`), and Word/LibreOffice refuse the file with no useful diagnostic.

The fixtures build that disagreement deterministically rather than hoping a
local pandoc reproduces it: the output package is renumbered exactly the way
the issue documents. What is under test is the remap, not pandoc's numbering.

Pure stdlib + python-docx; no pandoc, no docling, no Docker. Run:

    python -m pytest packages/document-converter/engine/tests -q
"""
import re
import shutil
import sys
import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK

ENGINE_DIR = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ENGINE_DIR))

from document_converter.cover_page import copy_template_cover_page  # noqa: E402

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
R = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
RELS_PART = "word/_rels/document.xml.rels"
DOC_PART = "word/document.xml"

HDR_FTR = [
    ("header", "default"),
    ("footer", "default"),
    ("first_page_header", "first"),
    ("first_page_footer", "first"),
    ("even_page_header", "even"),
    ("even_page_footer", "even"),
]


# --- fixtures --------------------------------------------------------------


def _make_template(path: Path) -> Path:
    """A template whose cover is its own section: six header/footer parts, and
    a paragraph-level `<w:sectPr>` that references all six."""
    d = Document()
    d.settings.odd_and_even_pages_header_footer = True
    section = d.sections[0]
    section.different_first_page_header_footer = True
    for attr, kind in HDR_FTR:
        getattr(section, attr).paragraphs[0].text = f"COVER-{attr}"

    d.add_paragraph("COVER TITLE")
    d.add_section(WD_SECTION.NEW_PAGE)
    d.add_paragraph("BODY")

    # The page break must live in the paragraph that CARRIES the cover section's
    # sectPr: `copy_template_cover_page` stops at the first page break, so a
    # break in an earlier paragraph would leave the sectPr uncopied and every
    # assertion below vacuous.
    for p in d.paragraphs:
        if p._element.find(f"{W}pPr/{W}sectPr") is not None:
            p.add_run().add_break(WD_BREAK.PAGE)
            break
    else:  # pragma: no cover - fixture guard
        raise AssertionError("fixture built no paragraph-level sectPr")

    d.save(path)
    return path


def _renumber_header_footer_rels(path: Path, offset: int) -> None:
    """Move the header/footer rIds of a package to another numeric range,
    keeping the package self-consistent.

    This is what pandoc does to the reporter's template as a side effect of
    rebuilding the rels: same parts, different ids. The offset makes the
    template's own ids UNRESOLVABLE here — the dangling `rId15` of the report.
    """
    with zipfile.ZipFile(path) as zf:
        parts = {n: zf.read(n) for n in zf.namelist()}

    rels = parts[RELS_PART].decode("utf-8")
    ids = re.findall(r'Id="(rId\d+)"[^>]*Target="((?:header|footer)\d+\.xml)"', rels)
    assert ids, "fixture has no header/footer relationships"

    old = [rid for rid, _ in ids]
    taken = set(re.findall(r'Id="(rId\d+)"', rels)) - set(old)
    mapping = {rid: f"rId{int(rid[3:]) + offset}" for rid in old}
    assert not (set(mapping.values()) & taken), "offset collides with an existing id"

    # Two-phase so a renumber never collides with an id it is about to write.
    def apply(text: str) -> str:
        for i, rid in enumerate(old):
            text = text.replace(f'"{rid}"', f'"__TMP{i}__"')
        for i, rid in enumerate(old):
            text = text.replace(f'"__TMP{i}__"', f'"{mapping[rid]}"')
        return text

    parts[RELS_PART] = apply(rels).encode("utf-8")
    parts[DOC_PART] = apply(parts[DOC_PART].decode("utf-8")).encode("utf-8")

    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, blob in parts.items():
            zf.writestr(name, blob)


def _make_output(path: Path, template: Path, offset: int = 10) -> Path:
    """Stand in for pandoc's `--reference-doc` output: the template's header and
    footer parts, a fresh body, and its own rel numbering."""
    shutil.copy(template, path)
    d = Document(path)
    for p in list(d.paragraphs):
        p._element.getparent().remove(p._element)
    d.add_paragraph("GENERATED BODY")
    d.save(path)
    if offset:
        _renumber_header_footer_rels(path, offset)
    return path


# --- helpers ---------------------------------------------------------------


def _rels(path: Path) -> dict:
    with zipfile.ZipFile(path) as zf:
        rels = zf.read(RELS_PART).decode("utf-8")
    return dict(re.findall(r'Id="(rId\d+)"[^>]*Target="([^"]+)"', rels))


def _referenced_ids(path: Path) -> set:
    with zipfile.ZipFile(path) as zf:
        doc = zf.read(DOC_PART).decode("utf-8")
    return set(re.findall(r'r:(?:id|embed|link)="(rId\d+)"', doc))


def _cover_sect_refs(path: Path) -> dict:
    """`{(kind, type): target}` for the copied cover section, resolved through
    the output's own rels — i.e. what Word will actually load."""
    rels = _rels(path)
    doc = Document(path)
    for sect in doc.element.body.iter(f"{W}sectPr"):
        if sect.getparent().tag != f"{W}pPr":
            continue  # body-final sectPr belongs to the generated body
        out = {}
        for ref in sect:
            if ref.tag in (f"{W}headerReference", f"{W}footerReference"):
                kind = ref.tag.replace(W, "").replace("Reference", "")
                out[(kind, ref.get(f"{W}type"))] = rels.get(ref.get(f"{R}id"))
        return out
    return {}


# --- the reported bug ------------------------------------------------------


def test_no_dangling_relationship_id_after_copy(tmp_path):
    """#506: every r:id in document.xml must resolve in document.xml.rels."""
    template = _make_template(tmp_path / "template.docx")
    output = _make_output(tmp_path / "out.docx", template)

    before = _referenced_ids(output)
    copy_template_cover_page(output, template)
    after = _referenced_ids(output)

    # Anti-vacuity: the copied cover must actually contribute references, else
    # "no dangling ids" would be true of a document that copied nothing.
    assert after - before, "the copied cover contributed no relationship ids"

    dangling = after - set(_rels(output))
    assert dangling == set(), f"dangling relationship ids: {sorted(dangling)}"


def test_cover_references_keep_their_meaning(tmp_path):
    """A remap by id alone would silently repoint `type='first'` at another
    part. Every reference must resolve to the SAME target as in the template."""
    template = _make_template(tmp_path / "template.docx")
    output = _make_output(tmp_path / "out.docx", template)

    expected = _cover_sect_refs(template)
    assert expected, "fixture lost its paragraph-level cover sectPr"

    copy_template_cover_page(output, template)

    assert _cover_sect_refs(output) == expected


def test_unshifted_package_is_unchanged(tmp_path):
    """When both numberings already agree the remap must be a no-op."""
    template = _make_template(tmp_path / "template.docx")
    output = _make_output(tmp_path / "out.docx", template, offset=0)

    copy_template_cover_page(output, template)

    assert _cover_sect_refs(output) == _cover_sect_refs(template)
    assert _referenced_ids(output) <= set(_rels(output))


def test_unresolvable_header_reference_is_dropped_not_dangling(tmp_path):
    """A header the output package simply does not have: drop the reference (a
    missing header is benign) rather than emit an id Word cannot resolve."""
    template = _make_template(tmp_path / "template.docx")
    output = _make_output(tmp_path / "out.docx", template)

    # Strip every even-page part from the output package.
    with zipfile.ZipFile(output) as zf:
        parts = {n: zf.read(n) for n in zf.namelist()}
    doomed = [rid for rid, tgt in _rels(output).items() if tgt in ("header3.xml", "footer3.xml")]
    rels = parts[RELS_PART].decode("utf-8")
    for rid in doomed:
        rels = re.sub(rf'<Relationship Id="{rid}"[^>]*/>', "", rels)
    parts[RELS_PART] = rels.encode("utf-8")
    for name in ("word/header3.xml", "word/footer3.xml"):
        parts.pop(name, None)
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, blob in parts.items():
            zf.writestr(name, blob)

    copy_template_cover_page(output, template)

    assert _referenced_ids(output) - set(_rels(output)) == set()
    refs = _cover_sect_refs(output)
    assert ("header", "even") not in refs and ("footer", "even") not in refs
    assert refs[("header", "first")] == "header2.xml"


def test_dangling_non_section_reference_fails_loudly(tmp_path):
    """Anything else unresolvable is a corrupt package: raise instead of writing
    a file Word will reject with 'source file could not be loaded'."""
    template = _make_template(tmp_path / "template.docx")

    # A cover-local reference that no rel backs, and that is not a header/footer.
    doc = Document(template)
    doc.paragraphs[0].runs[0]._element.set(f"{R}embed", "rId999")
    doc.save(template)

    output = _make_output(tmp_path / "out.docx", template)

    with pytest.raises(RuntimeError, match=r"rId999"):
        copy_template_cover_page(output, template)


def test_pre_existing_dangling_reference_in_the_body_is_caught(tmp_path):
    """The guard runs over the whole document before writing, not just over the
    elements this function copied — the output must never be written broken."""
    template = _make_template(tmp_path / "template.docx")
    output = _make_output(tmp_path / "out.docx", template)

    doc = Document(output)
    doc.paragraphs[0].add_run("x")._element.set(f"{R}embed", "rId777")
    doc.save(output)

    with pytest.raises(RuntimeError, match=r"rId777"):
        copy_template_cover_page(output, template)

    # …and the broken package was NOT written over.
    assert "rId777" in _referenced_ids(output)
